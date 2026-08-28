"""Build the HAMI (Hampi) source docx used by the chunker."""
from docx import Document

sections = [
    ("Introduction",
     "Hampi is a UNESCO World Heritage Site located in the Vijayanagara district of "
     "Karnataka, India. It stands on the banks of the Tungabhadra River and was the "
     "capital of the Vijayanagara Empire from 1336 to 1565. At its peak, Hampi was one "
     "of the largest and richest cities in the world, admired by Persian, Portuguese "
     "and Italian travellers for its wealth, temples and markets. Today the Hampi "
     "ruins spread across roughly 26 square kilometres and contain more than 1,600 "
     "surviving monuments, making it one of the largest open-air museum sites in the "
     "world. UNESCO inscribed Hampi on the World Heritage List in 1986 as the Group of "
     "Monuments at Hampi."),
    ("Founding of Vijayanagara",
     "The Vijayanagara Empire was founded in 1336 by two brothers, Harihara I and "
     "Bukka Raya I of the Sangama dynasty, on the banks of the Tungabhadra River. "
     "Legend recorded by the poet Nuniz says that the site was blessed by the sage "
     "Vidyaranya, after whom the capital was also called Vidyanagara. The city grew "
     "rapidly because of its defensive granite hills, its position on trade routes "
     "between the Deccan and the coast, and the sacred landscape of Pampa-kshetra "
     "associated with the goddess Pampa, a form of Parvati. The empire reached its "
     "greatest extent under Krishnadevaraya of the Tuluva dynasty, who ruled from 1509 "
     "to 1529 and expanded Vijayanagara control over most of South India."),
    ("Virupaksha Temple",
     "The Virupaksha Temple is the oldest and principal temple of Hampi and remains an "
     "active place of worship today. It is dedicated to Lord Shiva, worshipped here as "
     "Virupaksha, the patron deity of the Vijayanagara rulers. The temple predates the "
     "empire itself, with origins traced to the 7th century under the Chalukyas and "
     "Hoysalas. Its eastern gopuram, or gateway tower, rises about nine storeys and "
     "dominates the Hampi bazaar street. Inside are halls with finely carved pillars, "
     "shrines to Pampa and Bhuvaneshwari, and a courtyard where the annual chariot "
     "festival is held. The temple was greatly expanded by Krishnadevaraya in the "
     "early 16th century."),
    ("Vittala Temple and Stone Chariot",
     "The Vittala Temple is the most famous monument of Hampi and the finest example "
     "of Vijayanagara architecture. It is dedicated to Vittala, a form of Lord Vishnu. "
     "Its Maha Mantapa, or great hall, contains fifty-six musical pillars, known as "
     "the SaReGaMa pillars, which produce musical tones when tapped. The temple "
     "complex also contains the iconic Stone Chariot, a monolithic granite shrine "
     "modelled on a wooden temple chariot, which appears on the Indian fifty rupee "
     "note. In front of the temple lies the Hampi bazaar and a broad chariot street "
     "paved with granite slabs. Construction of the temple began in the 15th century "
     "and continued under successive rulers."),
    ("Hazararama Temple",
     "The Hazararama Temple, meaning the temple of a thousand Ramas, was built in the "
     "early 15th century by Devaraya I. It served as the private chapel of the royal "
     "family inside the royal enclosure, unlike the larger public temples. Its walls "
     "carry intricate bas-relief panels depicting scenes from the Ramayana in three "
     "horizontal bands that spiral around the outer walls. The temple also contains a "
     "well-preserved shrine and carved pillars showing Vishnu as Varaha and other "
     "avatars."),
    ("Lotus Mahal and Zenana Enclosure",
     "The Lotus Mahal is a graceful two-storeyed pavilion inside the Zenana Enclosure, "
     "the walled area reserved for the royal women of the court. The building blends "
     "Hindu and Islamic architectural styles, with lobed arched openings and a "
     "pyramidal tower that resembles a lotus bud. Its walls are plastered and its "
     "design kept the interior cool in summer. The Zenana Enclosure also contains the "
     "three watchtowers and the Elephant Stables, a long row of domed chambers built "
     "to house the royal elephants. The Elephant Stables are among the best-preserved "
     "buildings at Hampi and show a strong Indo-Islamic style with domes and arched "
     "ventilators."),
    ("Royal Enclosure and Mahanavami Dibba",
     "The Royal Enclosure was the ceremonial and administrative heart of the "
     "Vijayanagara capital. It contains the Mahanavami Dibba, a massive stepped stone "
     "platform about twelve metres high, from which the kings watched the Mahanavami "
     "festival, parades and wrestling matches. The enclosure also includes the remains "
     "of the king's palace, underground chambers, an audience hall with hundreds of "
     "pillar bases, and the stepped Pushkarani tank used for ritual bathing. An "
     "aqueduct system channelled water into the tank from the Kamalapuram lake."),
    ("Water and Irrigation Systems",
     "Hampi's rulers built an elaborate hydraulic system of canals, aqueducts, tanks "
     "and wells to support a large population in a semi-arid region. The Tungabhadra "
     "River was tapped by stone-lined canals, some carved through solid rock, that "
     "carried water to the city and its orchards. Large reservoirs such as the "
     "Kamalapuram tank and the Turthu canal system supplied agriculture and the royal "
     "baths. Aqueducts fed the stepped Pushkarani tank inside the Royal Enclosure. "
     "This water management system is considered one of the great engineering "
     "achievements of medieval South India."),
    ("Hampi Bazaar and Trade",
     "The Hampi Bazaar, also called Virupaksha Bazaar, is a nearly one kilometre long "
     "street running east from the Virupaksha Temple. Foreign travellers such as "
     "Abdur Razzaq, Domingo Paes and Fernao Nuniz described Hampi as a city of "
     "extraordinary wealth, with streets crowded with merchants selling diamonds, "
     "pearls, silk, spices and horses. Paes wrote that the city was the best provided "
     "city in the world. Trade connected the empire to Portuguese Goa, Persia and "
     "Southeast Asia, and horses were imported in large numbers for the Vijayanagara "
     "cavalry."),
    ("Decline of Hampi",
     "After the death of Krishnadevaraya in 1529 the empire weakened under his "
     "successors. In 1565 the Deccan Sultanates of Bijapur, Ahmadnagar, Golconda, "
     "Bidar and Berar united to form a confederacy against Vijayanagara. At the Battle "
     "of Talikota, also called the Battle of Rakkasa-Tangadi, the Vijayanagara army "
     "was defeated and its ruler Aliya Rama Raya was killed. The victorious armies "
     "then sacked and burned Hampi for several months. The city was never rebuilt on "
     "its former scale, and the capital was shifted south to Penugonda. The ruins "
     "gradually became overgrown and were reused as building material by later "
     "settlers."),
    ("Rediscovery and Conservation",
     "British surveyors and archaeologists such as Colin Mackenzie documented the "
     "Hampi ruins in the early 19th century. Systematic conservation began in the 20th "
     "century under the Archaeological Survey of India. In 1986 UNESCO inscribed the "
     "Group of Monuments at Hampi as a World Heritage Site. A Purushottama technique "
     "of aniconic worship and continued conservation campaigns by the ASI and the "
     "Karnataka state department have stabilized many monuments. Today Hampi is a "
     "major tourist destination visited by travellers from around the world, and part "
     "of it lies within the Hampi World Heritage Area Management Authority's "
     "protected zone."),
    ("Geography of the Hampi Region",
     "Hampi lies in the Vijayanagara district of Karnataka at about 15.34 degrees "
     "north latitude and 76.46 degrees east longitude. The landscape is dominated by "
     "striking granite boulder hills formed over three billion years ago, some of the "
     "oldest exposed rock in India. The Tungabhadra River flows northeast through the "
     "site, and the city was built on its southern bank. The region has a dry "
     "tropical climate, with very hot summers, a moderate monsoon and pleasant "
     "winters. Banana, sugarcane and paddy are cultivated along the river using the "
     "ancient canal system."),
    ("Timeline of Hampi",
     "7th century: early shrines including Virupaksha are worshipped at Pampa-kshetra. "
     "1336: Harihara I and Bukka Raya I found the Vijayanagara Empire and build the "
     "capital on the Tungabhadra. 14th to 15th century: the Sangama and Saluva "
     "dynasties expand the city and its temples. 1509 to 1529: Krishnadevaraya rules "
     "at the empire's height and expands the Vittala and Virupaksha temples. 1565: "
     "defeat at the Battle of Talikota followed by the sack of Hampi. 19th century: "
     "the ruins are surveyed and documented by British archaeologists. 1986: UNESCO "
     "lists the Group of Monuments at Hampi as a World Heritage Site."),
]

document = Document()
document.add_heading("HAMI", level=1)

for heading, body in sections:
    document.add_heading(heading, level=2)
    document.add_paragraph(body)

document.save("data/documents/HAMI.docx")
print("Created data/documents/HAMI.docx")
